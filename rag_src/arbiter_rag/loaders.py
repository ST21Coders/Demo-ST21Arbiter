"""Document loaders — turn the raw corpus into text + metadata for ingestion.

Only used at INGEST time (notebooks, pipelines), never by the query-only agent, so
pypdf/pandas are imported lazily to keep the agent runtime dependency-light.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

_EFFECTIVE_RE = re.compile(r"Effective\s+(\d{4}-\d{2}-\d{2})")
# Repeated page footer / subtitle noise we don't want in the retrievable text.
_FOOTER_RE = re.compile(r"FICTIONAL SAMPLE|Page\s+\d+", re.IGNORECASE)


def load_pdf_text(path: str | Path) -> str:
    """Extract all text from a PDF (requires the `data`/`notebook` extra: pypdf)."""
    from pypdf import PdfReader  # lazy

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def _clean_policy_text(text: str) -> str:
    """Drop the repeated page-footer lines so they don't pollute chunks/embeddings."""
    kept = [ln for ln in text.splitlines() if not _FOOTER_RE.search(ln)]
    # Collapse the runs of blank lines left behind so paragraph splitting stays clean.
    cleaned = "\n".join(kept)
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()


def iter_hr_documents(pdf_dir: str | Path) -> list[dict[str, Any]]:
    """Load every HR-policy PDF into a record with text + filterable metadata.

    doc_id and policy_category come from the filename (e.g. HR-LEAVE-001_leave.pdf);
    title/effective_date are parsed from the document body. These map directly onto
    the S3 Vectors `hr-policies` index schema.
    """
    pdf_dir = Path(pdf_dir)
    docs: list[dict[str, Any]] = []
    for path in sorted(pdf_dir.glob("*.pdf")):
        stem = path.stem  # e.g. "HR-LEAVE-001_leave"
        doc_id, _, category = stem.partition("_")
        text = _clean_policy_text(load_pdf_text(path))
        # Title = first content line that isn't the "Document ... Category ..." subtitle.
        title = next(
            (
                ln.strip()
                for ln in text.splitlines()
                if ln.strip() and not ln.strip().startswith("Document ")
            ),
            doc_id,
        )
        eff = _EFFECTIVE_RE.search(text)
        docs.append(
            {
                "doc_id": doc_id,
                "title": title,
                "policy_category": category or "general",
                "effective_date": eff.group(1) if eff else "",
                "state": "US-ALL",          # policies apply to all US outlets in this demo
                "access_level": "employee",  # RBAC tag injected server-side on retrieval
                "source_file": path.name,
                "source_uri": f"file://{path}",
                "text": text,
            }
        )
    return docs


# --------------------------------------------------------------------------- #
# Generic unstructured loader (DocuSearch path): pdf / docx / txt / md / json.
# --------------------------------------------------------------------------- #
_TEXT_EXTS = {".txt", ".md"}
SUPPORTED_DOC_EXTS = {".pdf", ".docx", ".txt", ".md", ".json"}


def load_docx_text(path: str | Path) -> str:
    """Extract paragraph + table text from a Word .docx (requires python-docx)."""
    from docx import Document  # lazy

    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:  # keep tabular content in Word docs
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def load_json_text(path: str | Path) -> str:
    """Flatten a JSON document to readable `dotted.key: value` lines for embedding."""
    import json  # local import keeps the module symmetrical with the lazy loaders

    data = json.loads(Path(path).read_text())
    lines: list[str] = []

    def _walk(prefix: str, node: Any) -> None:
        if isinstance(node, dict):
            for k, v in node.items():
                _walk(f"{prefix}.{k}" if prefix else str(k), v)
        elif isinstance(node, list):
            for i, v in enumerate(node):
                _walk(f"{prefix}[{i}]", v)
        else:
            lines.append(f"{prefix}: {node}" if prefix else str(node))

    _walk("", data)
    return "\n".join(lines).strip()


def load_document_text(path: str | Path) -> str:
    """Extract plain text from any SUPPORTED_DOC_EXTS file (dispatch by extension)."""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return load_pdf_text(path)
    if ext == ".docx":
        return load_docx_text(path)
    if ext == ".json":
        return load_json_text(path)
    if ext in _TEXT_EXTS:
        return Path(path).read_text(errors="ignore").strip()
    raise ValueError(f"Unsupported document type {ext!r}: {path}")


def iter_documents(folder: str | Path, *, recursive: bool = True) -> list[dict[str, Any]]:
    """Load every supported document under `folder` into a text + metadata record.

    Handles .pdf / .docx / .txt / .md / .json — the generic unstructured loader for the
    DocuSearch ingest path. doc_id/title default to the filename; callers needing richer
    domain metadata (e.g. HR policy_category) use iter_hr_documents instead.
    """
    folder = Path(folder)
    walk = folder.rglob("*") if recursive else folder.glob("*")
    paths = sorted(p for p in walk if p.is_file() and p.suffix.lower() in SUPPORTED_DOC_EXTS)
    docs: list[dict[str, Any]] = []
    for path in paths:
        text = load_document_text(path)
        if not text:
            continue
        title = next((ln.strip() for ln in text.splitlines() if ln.strip()), path.stem)[:200]
        docs.append(
            {
                "doc_id": path.stem,
                "title": title,
                "doc_type": path.suffix.lower().lstrip("."),
                "source_file": path.name,
                "source_uri": f"file://{path}",
                "text": text,
            }
        )
    return docs


def load_sales_dataframe(path: str | Path, sheet: str = "Sales"):
    """Load a sheet of the sales workbook into a pandas DataFrame (requires pandas)."""
    import pandas as pd  # lazy

    return pd.read_excel(path, sheet_name=sheet)


# Expected column contract of the Hawaiian-electronics sales CSVs (one row per transaction).
HAWAII_SALES_COLUMNS = [
    "transaction_id", "branch_id", "city", "island", "state", "timestamp", "sku",
    "product_description", "category", "quantity", "unit_cost", "unit_price", "revenue",
    "cost", "gross_margin", "salesperson", "customer_type", "payment_method",
    "inventory_remaining",
]
_NUMERIC_COLS = [
    "quantity", "unit_cost", "unit_price", "revenue", "cost", "gross_margin",
    "inventory_remaining",
]


def load_hawaii_sales(csv_dir: str | Path):
    """Concatenate the per-branch Hawaii sales CSVs into one typed DataFrame.

    `csv_dir` holds files like `HI001_Honolulu_HI_sales_2026-06-30.csv` — one per branch,
    all sharing HAWAII_SALES_COLUMNS. Numeric columns are coerced so downstream
    aggregation (serialization + the pandas ground-truth used by eval) is exact. Requires
    pandas (ingest-time only; the query agent never imports this module).
    """
    import pandas as pd  # lazy

    csv_dir = Path(csv_dir)
    paths = sorted(csv_dir.glob("*.csv"))
    if not paths:
        raise FileNotFoundError(f"No CSV files found in {csv_dir}")
    frames = [pd.read_csv(p) for p in paths]
    df = pd.concat(frames, ignore_index=True)

    missing = [c for c in HAWAII_SALES_COLUMNS if c not in df.columns]
    if missing:
        raise ValueError(f"Sales CSVs missing expected columns: {missing}")
    for col in _NUMERIC_COLS:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    return df


# --------------------------------------------------------------------------- #
# Generic tabular loader (Structured Analytics path): csv / excel / parquet / flat.
# --------------------------------------------------------------------------- #
TABULAR_EXTS = {".csv", ".xlsx", ".xls", ".parquet"}


def load_tabular(path_or_dir: str | Path, *, sheet: int | str = 0):
    """Load csv/excel/parquet — a single file or a directory of them — into one DataFrame.

    Concatenates every supported tabular file under a directory (schema inferred). Used by
    the Structured Analytics ingest path for arbitrary datasets; the Hawaii demo keeps
    load_hawaii_sales for its fixed column contract. Requires pandas (+ openpyxl for xlsx,
    pyarrow for parquet), all in the `data` extra — ingest-time only.
    """
    import pandas as pd  # lazy

    p = Path(path_or_dir)
    if p.is_file():
        paths = [p]
    else:
        paths = sorted(x for x in p.rglob("*") if x.is_file() and x.suffix.lower() in TABULAR_EXTS)
    if not paths:
        raise FileNotFoundError(f"No csv/excel/parquet files found in {path_or_dir}")

    frames = []
    for fp in paths:
        ext = fp.suffix.lower()
        if ext == ".csv":
            frames.append(pd.read_csv(fp))
        elif ext in (".xlsx", ".xls"):
            frames.append(pd.read_excel(fp, sheet_name=sheet))
        elif ext == ".parquet":
            frames.append(pd.read_parquet(fp))
    return pd.concat(frames, ignore_index=True) if len(frames) > 1 else frames[0]
